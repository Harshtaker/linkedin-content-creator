import datetime
import random
import csv
from docx import Document

# -------------------- BEST POSTING TIME --------------------
def best_posting_time(profile_type: str) -> str:
    """
    Suggests best posting time based on profile type.
    """
    times = {
        "Professional": "8–9 AM or 6–7 PM (Tuesday–Thursday)",
        "Student": "12–2 PM or 8–9 PM (Monday–Friday)",
        "Entrepreneur": "6–8 AM (Weekdays) or 10–11 AM (Saturday)",
        "Marketer": "9–11 AM (Tuesday–Friday)",
        "Developer": "10–11 AM or 4–6 PM (Weekdays)"
    }
    return times.get(profile_type, "9–11 AM (Weekdays)")

# -------------------- WEEKLY SCHEDULER --------------------
def generate_weekly_schedule(topics):
    """
    Generates a 7-day schedule based on the given list of topics.
    Returns a list of dictionaries with date, topic, and best posting time.
    """
    today = datetime.date.today()
    schedule = []

    for i, topic in enumerate(topics[:7]):
        post_date = today + datetime.timedelta(days=i)
        best_time = random.choice([
            "8–9 AM", "9–11 AM", "12–2 PM", "4–6 PM", "6–8 PM"
        ])
        schedule.append({
            "date": post_date.strftime("%A, %d %B %Y"),
            "topic": topic,
            "best_time": best_time
        })

    return schedule

# -------------------- EXPORT TO CSV --------------------
def export_to_csv(schedule, filename="weekly_schedule.csv"):
    """
    Exports the generated schedule to a CSV file.
    """
    with open(filename, mode="w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Date", "Topic", "Best Time"])
        for entry in schedule:
            writer.writerow([entry["date"], entry["topic"], entry["best_time"]])
    return filename

# -------------------- EXPORT TO DOCX --------------------
def export_to_docx(schedule, filename="weekly_schedule.docx"):
    """
    Exports the generated schedule to a DOCX file.
    """
    doc = Document()
    doc.add_heading("📅 Weekly LinkedIn Content Schedule", level=1)

    for entry in schedule:
        doc.add_paragraph(f"📅 {entry['date']}", style="List Bullet")
        doc.add_paragraph(f"🧠 Topic: {entry['topic']}")
        doc.add_paragraph(f"⏰ Best Time: {entry['best_time']}")
        doc.add_paragraph("")  # spacing

    doc.save(filename)
    return filename
