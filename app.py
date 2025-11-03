import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import datetime
from utils import generator, hashtags, image_fetch, content_planner as calendar


# --------------------- PAGE CONFIG ---------------------
st.set_page_config(
    page_title="LinkedIn Content Creator",
    page_icon="💼",
    layout="wide"
)

# --------------------- STYLES ---------------------
st.markdown("""
    <style>
        .main-title { 
            font-size: 42px; 
            font-weight: 700; 
            text-align: center; 
            color: #2E86C1;
            margin-bottom: 5px;
        }
        .subtitle {
            text-align: center;
            font-size: 18px;
            color: #7D7D7D;
            margin-bottom: 40px;
        }
        .stButton > button {
            width: 100%;
            background-color: #2E86C1 !important;
            color: white !important;
            font-size: 18px;
            border-radius: 10px !important;
        }
    </style>
""", unsafe_allow_html=True)

st.markdown("<p class='main-title'>💼 LinkedIn Content Creator</p>", unsafe_allow_html=True)
st.markdown("<p class='subtitle'>Generate high-quality posts and schedule your week easily.</p>", unsafe_allow_html=True)

# --------------------- NAVIGATION ---------------------
tabs = st.tabs(["📝 Generate Post", "📅 Content Scheduler"])

# ========================================================
# TAB 1: POST GENERATOR
# ========================================================
with tabs[0]:
    st.sidebar.header("🧠 Customize Your Post")
    topic = st.sidebar.text_input("Enter your topic:", placeholder="e.g., Power of Networking")
    profile = st.sidebar.selectbox("Profile Type", ["Professional", "Student", "Entrepreneur", "Marketer", "Developer"])
    tone = st.sidebar.selectbox("Tone", ["Professional", "Friendly", "Motivational", "Inspirational", "Educational"])
    length = st.sidebar.radio("Post Length", ["Short", "Medium", "Long"], index=1)

    st.sidebar.markdown("---")
    generate_button = st.sidebar.button("🚀 Generate Post")

    if generate_button:
        if not topic.strip():
            st.warning("⚠️ Please enter a topic to generate your post.")
        else:
            with st.spinner("✍️ Generating your LinkedIn post..."):
                post_text = generator.generate_post(topic, profile, tone, length)
                tags = hashtags.generate_hashtags(topic)
                image_url = image_fetch.fetch_image_url(topic)

            st.markdown("### 📝 Generated LinkedIn Post")
            st.success(post_text)

            st.markdown("### 🔖 Suggested Hashtags")
            st.write(" ".join([f"#{tag}" for tag in tags[:8]]))

            if image_url:
                st.markdown("### 🖼️ Image Suggestion")
                st.image(image_url, caption="Generated Image")
            else:
                st.info("No image available for this topic right now.")

            st.markdown("---")
            st.markdown("### ⏰ Best Posting Time Suggestion")
            best_time = calendar.best_posting_time(profile)
            st.info(f"Best time to post for a {profile}: **{best_time}**")
    else:
        st.info("👈 Enter details in the sidebar and click 'Generate Post' to begin!")

# ========================================================
# TAB 2: CONTENT SCHEDULER (New)
# ========================================================
with tabs[1]:
    st.header("📅 LinkedIn Post Scheduler")
    st.write("Plan and download your upcoming LinkedIn posts with auto-generated content!")

    # Initialize session state
    if "schedule" not in st.session_state:
        st.session_state["schedule"] = []

    topic = st.text_input("Enter post topic:", placeholder="e.g., Future of AI in 2025")
    date = st.date_input("Select posting date:", datetime.today())
    time = st.time_input("Select posting time:", datetime.now().time().replace(second=0, microsecond=0))
    auto_generate = st.checkbox("✨ Auto-generate post and hashtags", value=True)

    if st.button("➕ Add to Schedule"):
        if not topic.strip():
            st.warning("Please enter a topic before adding.")
        else:
            with st.spinner("Adding to your schedule..."):
                post_text, hashtags_list, image_url = "", "", ""
                if auto_generate:
                    try:
                        post_text = generator.generate_post(topic)
                        hashtags_list = " ".join([f"#{tag}" for tag in hashtags.generate_hashtags(topic)[:8]])
                        image_url = image_fetch.fetch_image_url(topic)
                    except Exception as e:
                        st.error(f"Auto-generation failed: {e}")

                st.session_state["schedule"].append({
                    "Topic": topic,
                    "Date": str(date),
                    "Time": str(time)[:5],
                    "Generated Post": post_text,
                    "Hashtags": hashtags_list,
                    "Image URL": image_url or ""
                })

                st.success(f"✅ Added: {topic} scheduled for {date} at {time.strftime('%H:%M')}")

    # Display current schedule
    if st.session_state["schedule"]:
        df = pd.DataFrame(st.session_state["schedule"])
        st.markdown("### 🗓️ Current Schedule")
        st.dataframe(df, use_container_width=True)

        # CSV Download
        csv = df.to_csv(index=False).encode("utf-8")
        st.download_button("⬇️ Download CSV", data=csv, file_name="LinkedIn_Schedule.csv", mime="text/csv")

        # DOCX Download
        doc = Document()
        doc.add_heading("LinkedIn Post Schedule", level=1)
        for idx, row in df.iterrows():
            doc.add_heading(f"Post {idx+1}: {row['Topic']}", level=2)
            if row["Generated Post"]:
                doc.add_paragraph(row["Generated Post"])
            if row["Hashtags"]:
                doc.add_paragraph(row["Hashtags"])
            doc.add_paragraph(f"Date: {row['Date']} | Time: {row['Time']}")
            if row["Image URL"]:
                doc.add_paragraph(f"Image URL: {row['Image URL']}")
            doc.add_paragraph("")

        doc_buf = BytesIO()
        doc.save(doc_buf)
        doc_buf.seek(0)
        st.download_button("⬇️ Download DOCX", data=doc_buf, file_name="LinkedIn_Schedule.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.info("No posts scheduled yet. Add one above!")

# ========================================================
# FOOTER
# ========================================================
st.markdown("""
    <hr>
    <div style='text-align: center; color: #A0A0A0; font-size: 14px;'>
        Built with ❤️ by <b>Harsh Shukla</b> | Powered by OpenRouter & Stable Diffusion APIs
    </div>
""", unsafe_allow_html=True)
